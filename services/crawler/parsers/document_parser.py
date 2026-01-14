"""
Document Parser - Parse Word/Excel documents

Extracts text from .docx and .xlsx files.
"""

from io import BytesIO
from typing import Dict, Any, Optional

from docx import Document
from openpyxl import load_workbook
from loguru import logger

from .base import BaseParser


class DocumentParser(BaseParser):
    """
    Parse Word and Excel documents.
    
    Supports:
    - .docx (Word)
    - .xlsx (Excel)
    """
    
    def parse(self, content: bytes, url: str) -> Optional[Dict[str, Any]]:
        """
        Parse document content.
        
        Args:
            content: Document bytes
            url: Source URL
            
        Returns:
            Parsed document
        """
        # Determine type from URL
        if url.lower().endswith(".docx"):
            return self._parse_docx(content, url)
        elif url.lower().endswith(".xlsx"):
            return self._parse_xlsx(content, url)
        else:
            logger.warning(f"Unknown document type: {url}")
            return None
    
    def _parse_docx(self, content: bytes, url: str) -> Optional[Dict[str, Any]]:
        """Parse Word document."""
        try:
            doc_file = BytesIO(content)
            doc = Document(doc_file)
            
            # Extract title (from first paragraph or filename)
            title = self._guess_title_from_url(url)
            if doc.paragraphs and doc.paragraphs[0].text:
                title = doc.paragraphs[0].text[:100]
            
            # Extract text
            text_parts = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            # Extract tables
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    text_parts.append(table_text)
            
            full_text = "\n\n".join(text_parts)
            
            # Create document
            return self._create_document(
                text=full_text,
                title=title,
                url=url,
                doc_type="docx",
                metadata={
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables),
                },
            )
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            return None
    
    def _parse_xlsx(self, content: bytes, url: str) -> Optional[Dict[str, Any]]:
        """Parse Excel spreadsheet."""
        try:
            xlsx_file = BytesIO(content)
            wb = load_workbook(xlsx_file, read_only=True, data_only=True)
            
            # Extract title from filename
            title = self._guess_title_from_url(url)
            
            # Extract text from all sheets
            text_parts = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                
                # Add sheet name as header
                text_parts.append(f"=== {sheet_name} ===")
                
                # Extract rows
                for row in sheet.iter_rows(values_only=True):
                    # Filter out None values
                    row_values = [str(cell) for cell in row if cell is not None]
                    if row_values:
                        text_parts.append(" | ".join(row_values))
            
            full_text = "\n".join(text_parts)
            
            # Create document
            return self._create_document(
                text=full_text,
                title=title,
                url=url,
                doc_type="xlsx",
                metadata={
                    "sheet_count": len(wb.sheetnames),
                    "sheet_names": wb.sheetnames,
                },
            )
            
        except Exception as e:
            logger.error(f"Error parsing XLSX: {e}")
            return None
    
    def _extract_table_text(self, table) -> str:
        """Extract text from Word table."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):  # Skip empty rows
                rows.append(" | ".join(cells))
        
        return "\n".join(rows)
    
    def _guess_title_from_url(self, url: str) -> str:
        """Guess title from URL filename."""
        parts = url.split("/")
        filename = parts[-1] if parts else "Untitled"
        
        # Remove extension
        if "." in filename:
            filename = ".".join(filename.split(".")[:-1])
        
        # Replace underscores/hyphens
        filename = filename.replace("_", " ").replace("-", " ")
        
        return filename.title()
